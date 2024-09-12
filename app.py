
from flask import Flask, request, render_template, jsonify , flash, redirect, url_for, session, logging ,send_file
from flask import redirect,url_for
import numpy as np
import pandas as pd
from keras.models import load_model
from datetime import datetime

# pdf 
from fpdf import FPDF  # PDF generation
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx2pdf import convert
from dotenv import load_dotenv
import os

# Load environment variables from .env file
load_dotenv()




import tensorflow as tf
import pickle


from flask_mysqldb import MySQL
from wtforms import Form, StringField, TextAreaField, PasswordField, validators
from passlib.hash import sha256_crypt
from functools import wraps
import MySQLdb.cursors

# data base
from flask_mysqldb import MySQL
import MySQLdb.cursors
import re
import requests
import bcrypt

# mail
from flask_mail import Mail, Message

# sms
import vonage
client = vonage.Client(key=os.getenv('SMS_KEY'), secret=os.getenv('SMS_SECRET'))
sms = vonage.Sms(client)


# s3 
import boto3
from botocore.exceptions import NoCredentialsError
# Initialize S3 client
s3 = boto3.client(
    's3',
    aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
    aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY_ID'),
    region_name=os.getenv('REGION_NAME')  # e.g., 'us-west-1'
)


# fill pdf 
from fillpdf import fillpdfs
form_fields = list(fillpdfs.get_form_fields('Final_Maintenance_Report.pdf').keys())


print(form_fields)

app = Flask(__name__)


# SQL
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'Dheeren'
app.config['MYSQL_DB'] = 'hackuser'

# init MYSQL
mysql = MySQL(app)


# import model
# model = load_model('Model7/ffnn3.keras')
# scaler_X = pickle.load(open('Model7/scaler_X.pkl', 'rb'))
# scaler_y = pickle.load(open('Model7/scaler_y.pkl', 'rb'))

model = load_model('Model8/ffnn4.keras')
scaler_X = pickle.load(open('Model8/scaler_X.pkl', 'rb'))
scaler_y = pickle.load(open('Model8/scaler_y.pkl', 'rb'))


# mail
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER')
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] =  os.getenv('MAIL_PASSWORD')
mail = Mail(app)

# model / Schema
class RegisterForm(Form):
    name = StringField('Name', [validators.Length(min=1, max=50)])
    username = StringField('Username', [validators.Length(min=4, max=25)])
    email = StringField('Email', [validators.Length(min=6, max=50)])
    password = PasswordField('Password', [
        validators.DataRequired(),
        validators.EqualTo('confirm', message='Passwords do not match')
    ])
    confirm = PasswordField('Confirm Password')




@app.route('/login', methods=['GET', 'POST'])
def login():
    msg = ''
    if request.method == 'POST' and 'username' in request.form and 'password' in request.form:
        username = request.form['username']
        password = request.form['password']
        
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM accounts WHERE username = %s', (username,))
        account = cursor.fetchone()
        
        if account and bcrypt.checkpw(password.encode('utf-8'), account['password'].encode('utf-8')):
            session['loggedin'] = True
            session['id'] = account['id']
            session['username'] = account['username']
            msg = 'Logged in successfully!'
            return redirect(url_for('home'))  # Redirect to a protected route upon successful login
        else:
            msg = 'Incorrect username / password!'
    
    return render_template('login.html', msg=msg)



def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('loggedin'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/logout')
@login_required
def logout():
    session.pop('loggedin', False)
    session.pop('id', None)
    session.pop('username', None)
    return redirect(url_for('login'))



@app.route('/register', methods=['GET', 'POST'])
def register():
    msg = ''
    if request.method == 'POST' and 'username' in request.form and 'password' in request.form and 'email' in request.form:
        username = request.form['username']
        password = request.form['password']
        email = request.form['email']
        phone_number = request.form.get('phone_number', '')  # Optional field handling
        
        if phone_number and not phone_number.startswith('+91'):
            phone_number = '+91' + phone_number
        # Hash the password
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
        
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM accounts WHERE username = %s', (username,))
        account = cursor.fetchone()
        
        if account:
            msg = 'Account already exists!'
        elif not re.match(r'[^@]+@[^@]+\.[^@]+', email):
            msg = 'Invalid email address!'
        elif not re.match(r'[A-Za-z0-9]+', username):
            msg = 'Username must contain only characters and numbers!'
        elif not username or not password or not email:
            msg = 'Please fill out the form!'
        else:
            cursor.execute(
                'INSERT INTO accounts (username, password, email, phone_number) VALUES (%s, %s, %s, %s)',
                (username, hashed_password, email, phone_number)
            )
            mysql.connection.commit()
            msg = 'You have successfully registered!'
    elif request.method == 'POST':
        msg = 'Please fill out the form!'
    
    return render_template('login.html', msg=msg)


@app.route('/', methods=['GET'])
@login_required
def home():
    return render_template('home.html')

@app.route('/analysis', methods=['GET'])
@login_required
def analysis():
    return render_template('analysis.html')

@app.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@app.route('/analytics')
def analytics():
    return render_template('dashboard.html')

@app.route('/messages')
def messages():
    return render_template('messages.html')

@app.route('/collections')
def collections():
    return render_template('collections.html')

@app.route('/application')
def application():
    return render_template('application.html')

@app.route('/application/shared')
def shared():
    return render_template('shared.html')

@app.route('/form', methods=['GET'])
@login_required
def form():
    return render_template('form.html')

@app.route('/upload-excel', methods=['GET'])
@login_required
def excel_form():
    print("slosknfk")
    return render_template('demo.html')

@app.route("/info", methods=['GET', 'POST'])
@login_required
def info():
    return render_template('diseases.html')   


# @app.route('/predict', methods=['POST'])
# @login_required
# def predict():
#     try:
#         # Get input values from the form
#         frequency = float(request.form.get('frequency'))
#         amplitude = float(request.form.get('amplitude'))
#         temperature = float(request.form.get('temperature'))
#         operating_hours = float(request.form.get('OperatingHours'))
#         print(frequency, amplitude, temperature, operating_hours)

#         # Convert input to a DataFrame with column names
#         input_data = pd.DataFrame([[amplitude, frequency, temperature, operating_hours]], 
#                                   columns=['Amplitude', 'Frequency', 'Temperature', 'OperatingHours'])

#         # Scale the input data using the loaded scaler
#         input_data_scaled = scaler_X.transform(input_data)

#         # Make prediction and inverse transform the result
#         prediction_scaled = model.predict(input_data_scaled)
#         predictions = scaler_y.inverse_transform(prediction_scaled)

#         # Convert predictions to native Python types
#         predicted_mass = float(predictions[0, 0])
#         predicted_lifespan = float(predictions[0, 1])
#         predicted_unbalance_force = float(predictions[0, 2])

#         # Define severity bins, labels, and map
#         bins = [0, 5, 10, 25, 40, 55, 65, np.inf]
#         labels = ['Negligible', 'Minor', 'Moderate', 'Significant', 'Serious', 'Severe', 'Critical']
#         severity_map = {'Negligible': 0, 'Minor': 1, 'Moderate': 2, 'Significant': 3, 'Serious': 4, 'Severe': 5, 'Critical': 6}

#         # Determine severity based on predicted mass
#         severity = pd.cut([predicted_mass], bins=bins, labels=labels, include_lowest=True).to_list()[0]
#         severity_numerical = severity_map[severity]

#         # Prepare data for rendering the template
#         result = {
#             'predicted_mass': predicted_mass,
#             'predicted_unbalance_force': predicted_unbalance_force,
#             'predicted_lifespan': predicted_lifespan,
#             'severity_name': severity,
#             'severity_numerical': severity_numerical
#         }

#         print(result)

#         user_id = session.get('id')
#         print(user_id)
#         # Check if the severity_numerical is 6 (Critical)
#         if severity_numerical == 6:
#             # Send email to the user
#             query = "SELECT phone_number, email FROM accounts WHERE id = %s"
#             cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
#             cursor.execute(query, (user_id,))
#             user = cursor.fetchone()

#             if user:
#                phone_number = user.get('phone_number')
#                email = user.get('email')
#                print(phone_number)  
#                print(email)
                
#             msg = Message(
#                 "Machine Maintenance Required",
#                 sender="derugaud@gmail.com",
#                 recipients=[email]
#             )
#             msg.body = "The machine needs maintenance immediately due to critical mass prediction."
            
#             # send_sms(phone_number)
#             # mail.send(msg)

#         return render_template('result.html', **result)

#     except Exception as e:
#         # Handle any exceptions that occur during processing
#         return f"Error during prediction: {str(e)}"


def save_prediction_results(prediction_data):
    fault_type = prediction_data.get('fault_type')
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    
    try:
        if fault_type == 'unbalance':
            data = prediction_data
            query = """
                INSERT INTO unbalance (predicted_mass, predicted_unbalance_force, predicted_lifespan,
                                       severity_name, radius, severity_numerical, frequency, amplitude,
                                       temperature, operating_hours)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            values = (
                data['predicted_mass'], data['predicted_unbalance_force'],
                data['predicted_lifespan'], data['severity_name'], data['radius'],
                data['severity_numerical'], data['frequency'], data['amplitude'],
                data['temperature'], data['operating_hours']
            )
            cursor.execute(query, values)
        
        elif fault_type == 'bearing':
            data = prediction_data
            query = """
                INSERT INTO bearing (predicted_load, predicted_vibration, predicted_wear,
                                     severity_name, severity_numerical, frequency, temperature, operating_hours)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """
            values = (
                data['predicted_load'], data['predicted_vibration'],
                data['predicted_wear'], data['severity_name'], data['severity_numerical'],
                data['frequency'], data['temperature'], data['operating_hours']
            )
            cursor.execute(query, values)
        
        elif fault_type == 'gear':
            data = prediction_data
            query = """
                INSERT INTO gear (predicted_torque, predicted_wear, severity_name,
                                  severity_numerical, frequency, temperature, operating_hours)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """
            values = (
                data['predicted_torque'], data['predicted_wear'],
                data['severity_name'], data['severity_numerical'], data['frequency'],
                data['temperature'], data['operating_hours']
            )
            cursor.execute(query, values)
        
        else:
            raise ValueError("Invalid fault type")

        mysql.connection.commit()
        fault_id = cursor.lastrowid  # Get the auto-generated ID
        cursor.close()
        return fault_id
    
    except Exception as e:
        cursor.close()
        raise e

def create_maintenance_report(prediction_data):
    machine_name = prediction_data.get('machine_name')
    severity = prediction_data.get('severity')
    report_link = prediction_data.get('report_link')
    fault_prediction_id = prediction_data.get('fault_prediction_id')
    comment = prediction_data.get('comment', None)
    maintenance_date = prediction_data.get('maintenance_date', None)
    
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    
    try:
        query = """
            INSERT INTO maintenance_reports (machine_name, severity, report_link, fault_prediction_id, comment, maintenance_date)
            VALUES (%s, %s, %s, %s, %s, %s)
        """
        values = (
            machine_name, severity, report_link, fault_prediction_id,
            comment, maintenance_date
        )
        cursor.execute(query, values)
        mysql.connection.commit()
        cursor.close()
    
    except Exception as e:
        cursor.close()
        raise e


def get_severity_name(predicted_severity_numerical):
    if 0 <= predicted_severity_numerical <= 0.5:
        return 'Negligible'
    elif 0.51 <= predicted_severity_numerical <= 1.5:
        return 'Minor'
    elif 1.6 <= predicted_severity_numerical <= 2.5:
        return 'Moderate'
    elif 2.6 <= predicted_severity_numerical <= 3.5:
        return 'Significant'
    elif 3.6 <= predicted_severity_numerical <= 4.5:
        return 'Serious'
    elif 4.6 <= predicted_severity_numerical <= 5.5:
        return 'Severe'
    else:  # 5.6 and above
        return 'Critical'



@app.route('/predict', methods=['POST'])
def predict():
    try:
        # Get input values from the form
        frequency = float(request.form.get('frequency'))
        amplitude = float(request.form.get('amplitude'))
        temperature = float(request.form.get('temperature'))
        operating_hours = float(request.form.get('OperatingHours'))
        radius = float(request.form.get('radius'))
        
        # machine_name = str(request.form.get('machine_name')) "
        machine_name = "m-1"

        # Convert input to a DataFrame with column names
        input_data = pd.DataFrame([[amplitude, frequency, temperature, operating_hours,radius]], 
                                  columns=['Amplitude', 'Frequency', 'Temperature', 'OperatingHours','Radius'])

        
        # Scale the input data using the loaded scaler
        input_data_scaled = scaler_X.transform(input_data)

        # Make prediction and inverse transform the result
        prediction_scaled = model.predict(input_data_scaled)
        predictions = scaler_y.inverse_transform(prediction_scaled)

        # Convert predictions to native Python types
        predicted_mass = float(predictions[0, 0])
        predicted_lifespan = float(predictions[0, 1])
        predicted_unbalance_force = float(predictions[0, 2])
        predicted_severity_numerical = float(predictions[0, 3])
        
        severity = get_severity_name(predicted_severity_numerical);
        result = {
            'fault_type':'unbalance',
            'predicted_mass': predicted_mass,
            'predicted_unbalance_force': predicted_unbalance_force,
            'predicted_lifespan': predicted_lifespan,
            'severity_name': severity,
            'radius': radius,
            'severity_numerical': predicted_severity_numerical  , # severity_numerical,
            'frequency': frequency,
            'amplitude': amplitude,
            'temperature': temperature,
            'operating_hours': operating_hours
        }
        print(result)
        

        user_id = session.get('id')
        # Generate PDF and get the file path
        
        pdf_file_path = fill_pdf(result)
        
        # fault_id = save_prediction_results(result)
        
        # create_maintenance_report({
        #     'machine_name': machine_name,
        #     'severity': severity,
        #     'report_link': "pdf_file_path",  # Placeholder for PDF path
        #     'fault_prediction_id': fault_id,
        #     'comment': None,
        #     'maintenance_date': None
        # })
        
        # Check if the severity_numerical is 6 (Critical)
        if predicted_severity_numerical >= 5.5:
            # Send email to the user
            query = "SELECT phone_number, email FROM accounts WHERE id = %s"
            cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
            cursor.execute(query, (user_id,))
            user = cursor.fetchone()

            if user:
                phone_number = user.get('phone_number')
                email = user.get('email')
                print(phone_number)
                
                send_maintenance_email(email,pdf_file_path)
                send_sms(phone_number,pdf_file_path)
        
        # Render the result template with a link to download the PDF
        return render_template('result.html', **result, pdf_file_path=pdf_file_path)
        # return render_template('result.html', **result, pdf_file_path="pdf_file_path")
    except Exception as e:
        # Handle any exceptions that occur during processing
        return f"Error during prediction: {str(e)}"

def send_sms(phone_number,pdf_file_path):
    print(phone_number)
    message_text = (
        "Maintenance Alert!\n\n"
        "The machine requires immediate maintenance due to a critical mass prediction.\n\n"
        "Download Report:\n"
        f"{pdf_file_path}\n\n"
        "Please address the issue promptly to avoid further complications. [FREE SMS DEMO, TEST MESSAGE]"
    )

    # Send the SMS message
    responseData = sms.send_message(
        {
            "from": "DAYNA4MITE",
            "to": phone_number,
            "text": message_text,
        }
    )

    if responseData["messages"][0]["status"] == "0":
      print("Message sent successfully.")
    else:
      print(f"Message failed with error: {responseData['messages'][0]['error-text']}")

def send_maintenance_email(email, pdf_file_path):
    # Create the email message
    msg = Message(
        "⚠️ Machine Maintenance Alert ⚠️",
        sender="derugaud@gmail.com",
        recipients=[email]
    )

    # Email body with structured content and clear instructions
    msg.body = (
        "Dear User,\n\n"
        "Our monitoring system has detected a critical issue that requires immediate attention:\n\n"
        "🚨 *Maintenance Required*: The machine has predicted a critical mass condition that needs urgent maintenance.\n\n"
        "📄 *Download Report*: To view the detailed analysis and recommended actions, please download the report using the link below:\n"
        f"{pdf_file_path}\n\n"
        "🔍 Review the report promptly to address the issue and prevent further complications.\n\n"
        "If you have any questions or need assistance, please contact our support team immediately.\n\n"
        "Thank you for your attention to this urgent matter.\n\n"
        "Best regards,\n"
        "Maintenance Team"
    )

    # HTML version of the email for more formatting and clickable link
    msg.html = (
        "<p>Dear User,</p>"
        "<p>Our monitoring system has detected a critical issue that requires immediate attention:</p>"
        "<ul>"
        "<li><strong>Maintenance Required:</strong> The machine has predicted a critical mass condition that needs urgent maintenance.</li>"
        "</ul>"
        "<p><strong>Download Report:</strong> To view the detailed analysis and recommended actions, please download the report using the link below:</p>"
        f"<p><a href='{pdf_file_path}'>Download Maintenance Report</a></p>"
        "<p>Review the report promptly to address the issue and prevent further complications.</p>"
        "<p>If you have any questions or need assistance, please contact our support team immediately.</p>"
        "<p>Thank you for your attention to this urgent matter.</p>"
        "<p>Best regards,<br>Maintenance Team</p>"
    )

    # Print message for debugging
    mail.send(msg)
    print(msg)

def generate_prediction_pdf(result):
    # Define the path to save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = f"/mnt/data/Predictive_Maintenance_Report_{timestamp}.docx"
    pdf_path = docx_path.replace(".docx", ".pdf")

    # Create a new document
    doc = Document()
    doc.add_heading('Fault Report: Predictive Maintenance for Rotating Machine', 0)

    # Machine Information
    doc.add_heading('Machine Information', level=1)
    doc.add_paragraph('Machine Name: Rotating Machine')
    doc.add_paragraph('Machine Type: Shaft-based Rotating Machine')
    doc.add_paragraph('Machine No.: RM-001')

    # Data Summary in a table format
    doc.add_heading('Data Summary', level=1)
    data_summary_table = doc.add_table(rows=1, cols=2)
    hdr_cells = data_summary_table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'

    # Adding borders to the table
    tbl = data_summary_table._tbl
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))
    tbl.tblPr.append(tblBorders)

    # Adding data summary values
    data_summary = [
        ('Predicted Mass', f'{result["predicted_mass"]} grams'),
        ('Predicted Unbalance Force', f'{result["predicted_unbalance_force"]} N'),
        ('Predicted Lifespan', f'{result["predicted_lifespan"]} hours'),
        ('Severity Index', f'{result["severity_numerical"]} ({result["severity_name"]})')
    ]
    for item, value in data_summary:
        row_cells = data_summary_table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = value

    # Fault Occurrence
    doc.add_heading('Fault Occurrence', level=1)
    doc.add_paragraph('Fault Type: Shaft Unbalance')
    doc.add_paragraph('Time of Fault: 2024-09-07 10:30 AM')
    doc.add_paragraph('Date of Fault: 2024-09-07')

    # Severity Mapping with color representation
    doc.add_heading('Severity Mapping', level=1)
    severity_table = doc.add_table(rows=1, cols=3)
    hdr_cells = severity_table.rows[0].cells
    hdr_cells[0].text = 'Predicted Mass (g)'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Meaning'

    severity_data = [
        ('0 - 4', '0', 'Negligible', RGBColor(40, 167, 69)),           # Green
        ('5 - 9', '1', 'Minor', RGBColor(111, 187, 111)),              # Light Green
        ('10 - 24', '2', 'Moderate', RGBColor(253, 216, 53)),          # Yellow
        ('25 - 39', '3', 'Significant', RGBColor(242, 158, 36)),       # Orange
        ('40 - 54', '4', 'Serious', RGBColor(243, 115, 33)),           # Dark Orange
        ('55 - 65', '5', 'Severe', RGBColor(240, 58, 23)),             # Red-Orange
        ('Above 65', '6', 'Critical', RGBColor(220, 53, 69))           # Red
    ]

    # Adding severity data with color-coded Meaning descriptions
    for mass, severity, meaning, color in severity_data:
        row_cells = severity_table.add_row().cells
        row_cells[0].text = mass
        row_cells[1].text = severity
        meaning_run = row_cells[2].paragraphs[0].add_run(meaning)
        meaning_run.font.color.rgb = color

    # Adding borders to the severity table
    tbl = severity_table._tbl
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>' % nsdecls('w'))
    tbl.tblPr.append(tblBorders)

    # Recommendations and Repair Suggestions
    doc.add_heading('Recommendations and Repair Suggestions', level=1)
    doc.add_paragraph('Recommendation: Reduce operating load and check alignment of the shaft. '
                      'Perform dynamic balancing to reduce unbalance force.')
    doc.add_paragraph('Repair Suggestion: Add or remove mass from the shaft to counteract the imbalance. '
                      'Regularly check the accelerometer readings and monitor temperature closely.')

    # Maintenance Schedule Based on Severity
    doc.add_heading('Maintenance Schedule Based on Severity', level=1)
    maintenance_points = [
        'Severity 0-1 (Negligible/Minor): Perform routine visual inspections every 3 months.',
        'Severity 2 (Moderate): Conduct detailed inspections every 2 months and check unbalance force.',
        'Severity 3 (Significant): Perform dynamic balancing every month.',
        'Severity 4 (Serious): Immediate inspection required. Dynamic balancing and shaft realignment within the week.',
        'Severity 5 (Severe): Immediate action required. Halt operation and replace the shaft if needed.',
        'Severity 6 (Critical): Immediate shutdown of the machine and full maintenance, including rotor replacement.'
    ]
    for point in maintenance_points:
        doc.add_paragraph(f'- {point}', style='ListBullet')

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(
        "The rotating machine exhibited unbalance due to uneven mass distribution on the shaft. "
        "Severity levels indicate potential issues ranging from minor to critical. Regular maintenance and balancing of "
        "the rotor are essential to prolong the machine's lifespan and ensure smooth operation. Immediate corrective measures "
        "should be taken to avoid further damage and minimize downtime."
    )
    
    # Save the document
    doc.save(docx_path)

    # Convert DOCX to PDF (Ensure convert function is correctly defined/imported)
    convert(docx_path, pdf_path)
    print(docx_path)
    
    
    try:
        bucket_name = 'hackwin'
        s3_file_name = f'reports/{os.path.basename(pdf_path)}'
        
        s3.upload_file(pdf_path, bucket_name, s3_file_name)
        
        # Generate a public URL for the file
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_file_name}"
        
        # Remove the local files after uploading
        os.remove(docx_path)
        os.remove(pdf_path)
        
        return s3_url
    
    except FileNotFoundError:
        return "The file was not found"
    
    except NoCredentialsError:
        return "Credentials not available"

    

    except Exception as e:
        print(f"Error during prediction: {e}")
        return f"Error occurred: {e}"

def fill_pdf(result):
    # Define the data dictionary to fill the PDF fields
    data_dict = {
    'Date': datetime.now().strftime('%Y-%m-%d'),  # Current date
    'Frequency': result.get('frequency', ""),  # Frequency from result
    'Amplitude': result.get('amplitude', ""),  # Amplitude from result
    'Radius': result.get('radius', ""),  # Radius from result
    'Mass': result.get('predicted_mass', ""),  # Predicted Mass from result
    'Unbalance Force': result.get('predicted_unbalance_force', ""),  # Predicted Unbalance Force from result
    'Temperature': result.get('temperature', ""),  # Temperature from result
    'Operating Hours': result.get('operating_hours', ""),  # Operating Hours from result
    'Severity Index': f"{result.get('severity_name', '')} {result.get('severity_numerical', '')}",  # Severity Index
    'Lifespan': result.get('predicted_lifespan', ""),  # Predicted Lifespan from result
    'Date of Occurance': datetime.now().strftime('%Y-%m-%d'),  # Current date for Date of Occurrence
    'Time of Occurance': datetime.now().strftime('%H:%M:%S'),  # Current time for Time of Occurrence
    'Remark': "",  # Leave blank or update as needed
    }
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name = f"Predictive_Maintenance_Report_{timestamp}.pdf"

    # Fill the PDF with data from the dictionary and create a new file
    fillpdfs.write_fillable_pdf('Final_Maintenance_Report.pdf', name, data_dict, flatten=True)

    # Define the S3 client
    
    try:
        # Upload the file to S3
        bucket_name = 'hackwin'
        s3_file_name = f'reports/{os.path.basename(name)}'
        s3.upload_file(name, bucket_name, s3_file_name)
        
        # Generate a public URL for the file
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_file_name}"
        
        # Remove the local file after uploading
        os.remove(name)
        
        return s3_url
    except FileNotFoundError:
        return "The file was not found"
    except NoCredentialsError:
        return "Credentials not available"
    except Exception as e:
        # Ensure the local file is deleted even if an error occurs
        if os.path.exists(name):
            os.remove(name)
        return f"An error occurred: {str(e)}"

@app.route('/download_pdf')
def download_pdf():
    pdf_file_path = request.args.get('pdf_file_path')
    if os.path.exists(pdf_file_path):
        return send_file(pdf_file_path, as_attachment=True)
    return "File not found."




if __name__ == '__main__':
    app.secret_key='secret123'
    app.run(debug=True,port=5000)